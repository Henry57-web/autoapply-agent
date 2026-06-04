from io import BytesIO

from docx import Document
from docx.shared import Pt


def build_resume_docx(title: str, content: str) -> BytesIO:
    document = _new_document()
    document.add_heading(title, level=0)
    _add_plain_text(document, content)
    return _serialize(document)


def build_cover_letter_docx(title: str, content: str) -> BytesIO:
    document = _new_document()
    document.add_heading(title, level=0)
    _add_plain_text(document, content)
    return _serialize(document)


def _new_document() -> Document:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    return document


def _add_plain_text(document: Document, content: str) -> None:
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
        elif stripped.startswith(("-", "•")):
            document.add_paragraph(stripped.lstrip("-• ").strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)


def _serialize(document: Document) -> BytesIO:
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
