import unittest

from docx import Document

from app.services.docx_export import build_cover_letter_docx, build_resume_docx


class DocxExportTests(unittest.TestCase):
    def test_build_resume_docx_preserves_resume_content(self) -> None:
        buffer = build_resume_docx("Tailored Resume - Data Engineer", "SUMMARY\n- Built ETL pipelines")
        document = Document(buffer)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]

        self.assertIn("Tailored Resume - Data Engineer", paragraphs)
        self.assertIn("SUMMARY", paragraphs)
        self.assertIn("Built ETL pipelines", paragraphs)

    def test_build_cover_letter_docx_preserves_cover_letter(self) -> None:
        buffer = build_cover_letter_docx("Cover Letter", "Dear Hiring Manager,\nI am interested.")
        document = Document(buffer)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]

        self.assertIn("Cover Letter", paragraphs)
        self.assertIn("Dear Hiring Manager,", paragraphs)
        self.assertIn("I am interested.", paragraphs)


if __name__ == "__main__":
    unittest.main()
